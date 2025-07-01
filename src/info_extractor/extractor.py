import sys
from pathlib import Path

# Ensure project root is in sys.path for absolute imports
project_root = Path(__file__).resolve().parent.parent.parent
if str(project_root) not in sys.path:
    sys.path.insert(0, str(project_root))

import os
import json
import re
import logging
import pandas as pd
from openai import OpenAI
from base64 import b64encode
from typing import Dict, Any, Optional, Union
import io
import hashlib
import docx

# Import PDF processing libraries
import pdfplumber
import camelot

# NEW: Import Google AI library for Gemini
import google.generativeai as genai
from google.generativeai.types import GenerateContentResponse, GenerationConfig

# Import project modules
from config import settings, prompts

logger = logging.getLogger(__name__)

class InfoExtractor:
    """
    支持PDF/Excel/图片/Word的工业参数提取系统
    - 新增了对 .docx 文件的支持
    - 新增了基于文件内容的缓存功能
    - 新增了使用 Gemini 原生API 直接处理PDF的现代化方法
    """
    def __init__(self):
        # General settings
        self.output_dir = settings.OUTPUT_DIR
        self.temperature = settings.LLM_TEMPERATURE
        self.request_timeout = settings.LLM_REQUEST_TIMEOUT

        # OpenAI-compatible client setup (for legacy/other files)
        self.openai_api_key = settings.LLM_API_KEY
        self.openai_model = settings.LLM_MODEL_NAME
        self.openai_api_url = settings.LLM_API_URL

        # Gemini client setup (for new PDF method)
        self.pdf_processing_method = settings.PDF_PROCESSING_METHOD
        self.gemini_api_key = settings.GEMINI_API_KEY
        self.gemini_model = settings.GEMINI_MODEL_NAME

        # Configure Gemini client if needed
        if self.pdf_processing_method == 'gemini':
            if not self.gemini_api_key:
                logger.error("PDF处理方法为 'gemini'，但未提供 GEMINI_API_KEY。")
            else:
                try:
                    genai.configure(api_key=self.gemini_api_key)
                    logger.info(f"Gemini 客户端配置成功，使用模型: {self.gemini_model}")
                except Exception as e:
                    logger.error(f"配置 Gemini 客户端失败: {e}")
        
        # Cache setup
        self.cache_dir = settings.OUTPUT_DIR / ".cache"
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        self.json_proc = self.JSONProc(self)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"初始化完成，输出目录：{self.output_dir}")

    def detect_file_type(self, file_path: Path) -> str:
        """精准文件类型检测 (已支持 .docx)"""
        ext = file_path.suffix.lower()
        if ext in ['.xls', '.xlsx', '.xlsm', '.xlsb', '.csv']: return "excel"
        if ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', '.tiff', '.tif']: return "image"
        if ext == '.pdf': return "pdf"
        if ext == '.docx': return "word"
        raise ValueError(f"不支持的文件类型：{ext}")

    def extract_parameters(self, file_path: Union[str, Path], output_filename: Optional[str] = None) -> Optional[Dict[str, Any]]:
        """增强型统一提取入口，根据配置分发任务"""
        file_path = Path(file_path)
        try:
            if not file_path.exists():
                raise FileNotFoundError(f"文件不存在：{file_path}")

            cache_path = self._get_cache_path(file_path)
            if cache_path.exists():
                logger.info(f"发现有效缓存，从 {cache_path} 加载结果。")
                with open(cache_path, 'r', encoding='utf-8') as f:
                    return json.load(f)

            file_type = self.detect_file_type(file_path)
            logger.info(f"开始处理 {file_type.upper()} 文件：{file_path.name} (无缓存)")

            api_response = None
            # --- API Call Dispatcher ---
            if file_type == 'pdf' and self.pdf_processing_method == 'gemini':
                logger.info("分发任务：使用 Gemini 原生方法处理 PDF。")
                api_response = self._execute_gemini_pdf_request(file_path)
            else:
                logger.info(f"分发任务：使用传统方法处理 {file_type}。")
                processed_data = self._process_file_content_legacy(file_path, file_type)
                if processed_data:
                    api_response = self._execute_openai_request(processed_data)
            
            if not api_response:
                logger.error("API 请求未能返回有效响应。")
                return None

            return self._handle_api_response(api_response, file_path, output_filename, cache_path)

        except Exception as e:
            logger.error(f"处理文件 {file_path.name} 失败：{str(e)}", exc_info=True)
            return None

    def _execute_gemini_pdf_request(self, file_path: Path) -> Optional[GenerateContentResponse]:
        """
        使用 Gemini File API 直接处理PDF文件。
        该方法会上传文件，生成内容，然后删除上传的文件。
        """
        uploaded_file = None
        try:
            if not self.gemini_api_key:
                logger.error("无法执行 Gemini 请求：API密钥未设置。")
                return None
            
            logger.info("正在准备 Gemini File API 请求...")
            
            # 1. Upload the file using the File API
            logger.info(f"正在上传文件到 Gemini File API: {file_path.name}")
            uploaded_file = genai.upload_file(path=file_path, display_name=file_path.name)
            logger.info(f"文件上传成功: {uploaded_file.uri}")

            # 2. Prepare the model and prompt for content generation
            model = genai.GenerativeModel(self.gemini_model)
            prompt_part = prompts.LLM_EXTRACTION_SYSTEM_PROMPT
            generation_config = GenerationConfig(response_mime_type="application/json")
            
            logger.info(f"向 Gemini ({self.gemini_model}) 发送请求...")
            response = model.generate_content(
                [uploaded_file, prompt_part],
                generation_config=generation_config
            )
            logger.info("已收到 Gemini API 响应。")
            return response

        except Exception as e:
            logger.error(f"执行 Gemini PDF 请求时出错: {e}", exc_info=True)
            return None
        finally:
            # 4. Clean up the uploaded file if it exists
            if uploaded_file:
                try:
                    logger.info(f"正在从 Gemini File API 删除文件: {uploaded_file.name}")
                    genai.delete_file(name=uploaded_file.name)
                    logger.info("文件删除成功。")
                except Exception as e_del:
                    logger.warning(f"删除上传的文件失败: {e_del}")


    def _execute_openai_request(self, processed_data: Dict) -> Optional[Any]:
        """使用 OpenAI 兼容的API处理预处理后的数据"""
        try:
            if not self.openai_api_key:
                logger.error("无法执行 OpenAI 请求：API密钥未设置。")
                return None

            client = OpenAI(api_key=self.openai_api_key, base_url=self.openai_api_url, timeout=self.request_timeout)
            messages = [{"role": "system", "content": prompts.LLM_EXTRACTION_SYSTEM_PROMPT}, {"role": "user", "content": []}]
            
            user_content = []
            if processed_data["mode"] == "image_url":
                user_content.append({"type": "image_url", "image_url": {"url": f"data:{processed_data['type']};base64,{processed_data['data']}"}})
            if "text_content" in processed_data:
                user_content.append({"type": "text", "text": processed_data["text_content"]})
            if "table_content_csv" in processed_data:
                 user_content.append({"type": "text", "text": f"参考表格数据 (CSV):\n{processed_data['table_content_csv']}"})

            messages[1]["content"] = user_content
            
            logger.info(f"向 OpenAI 兼容API ({self.openai_model}) 发送请求...")
            api_result = client.chat.completions.create(model=self.openai_model, messages=messages, temperature=self.temperature)
            logger.info("已收到 OpenAI 兼容API 响应。")
            return api_result
        except Exception as e:
            logger.error(f"执行 OpenAI 请求时出错: {e}", exc_info=True)
            return None

    def _process_file_content_legacy(self, file_path: Path, file_type: str) -> Optional[Dict]:
        """传统的文件内容预处理器 (用于非Gemini-PDF场景)"""
        try:
            # PDF (Legacy Method)
            if file_type == "pdf":
                # This part is now only called if PDF_PROCESSING_METHOD is 'legacy'
                image_content = self._encode_to_base64(file_path)
                text_content = "".join(page.extract_text() or "" for page in pdfplumber.open(file_path).pages)
                tables = camelot.read_pdf(str(file_path), flavor='lattice', pages='all')
                table_content_csv = "\n\n".join(f"--- Table {i+1} ---\n{table.df.to_csv(index=False)}" for i, table in enumerate(tables))
                return {"mode": "image_url", "data": image_content, "type": "application/pdf", "text_content": text_content, "table_content_csv": table_content_csv}
            
            # Excel
            elif file_type == "excel":
                df = pd.read_excel(file_path)
                return {"mode": "text", "text_content": df.to_csv(index=False)}
            
            # Image
            elif file_type == "image":
                return {"mode": "image_url", "data": self._encode_to_base64(file_path), "type": self._get_image_mime_type(file_path)}
            
            # Word
            elif file_type == "word":
                doc = docx.Document(file_path)
                text_content = [p.text for p in doc.paragraphs]
                for i, table in enumerate(doc.tables):
                    text_content.append(f"\n--- Table {i+1} ---")
                    for row in table.rows:
                        text_content.append(", ".join(cell.text.strip() for cell in row.cells))
                return {"mode": "text", "text_content": "\n".join(text_content)}

        except Exception as e:
            logger.error(f"内容预处理异常 ({file_type}): {e}", exc_info=True)
            return None

    def _handle_api_response(self, response: Any, file_path: Path, output_filename: Optional[str], cache_path: Path) -> Optional[Dict]:
        """统一处理来自不同API的响应"""
        raw_content = ""
        try:
            # Detect response type and extract content
            if isinstance(response, GenerateContentResponse):
                logger.debug("处理 Gemini 响应。")
                if not response.candidates:
                    logger.error("Gemini 响应中没有候选内容。")
                    if hasattr(response, 'prompt_feedback'):
                        logger.error(f"Prompt feedback: {response.prompt_feedback}")
                    return None
                raw_content = response.text
            elif hasattr(response, 'choices'): # OpenAI-like response
                logger.debug("处理 OpenAI 兼容响应。")
                if not response.choices:
                    logger.error("OpenAI 响应中 choices 为空。")
                    return None
                raw_content = response.choices[0].message.content
            else:
                logger.error(f"未知的API响应格式: {type(response)}")
                return None

            cleaned_json = self._clean_json_response(raw_content)
            result_dict = json.loads(cleaned_json)

            # Save main output file
            output_filename = output_filename or f"{file_path.stem}_analysis.json"
            output_path = self.output_dir / output_filename
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(result_dict, f, ensure_ascii=False, indent=2)
            logger.info(f"成功保存分析结果：{output_path}")

            # Write to cache
            self._write_to_cache(result_dict, cache_path)
            return result_dict

        except json.JSONDecodeError as e:
            logger.error(f"JSON解析失败：{e}")
            logger.debug(f"无法解析的原始响应内容：{raw_content[:500]}...")
            return None
        except Exception as e:
            logger.error(f"结果处理或保存时异常：{e}", exc_info=True)
            return None

    # ---------- 工具方法 ----------
    def _get_cache_path(self, file_path: Path) -> Path:
        hasher = hashlib.sha256()
        with open(file_path, 'rb') as f: hasher.update(f.read())
        return self.cache_dir / f"{hasher.hexdigest()}.json"

    def _write_to_cache(self, data: Dict, cache_path: Path):
        try:
            with open(cache_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            logger.info(f"结果已成功写入缓存: {cache_path}")
        except Exception as e:
            logger.warning(f"写入缓存文件失败: {e}")

    def _get_image_mime_type(self, file_path: Path) -> str:
        return {'.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.png': 'image/png', '.gif': 'image/gif', '.webp': 'image/webp'}.get(file_path.suffix.lower(), 'application/octet-stream')

    def _encode_to_base64(self, file_path: Path) -> str:
        with open(file_path, "rb") as f:
            return b64encode(f.read()).decode('utf-8')

    def _clean_json_response(self, response_str: str) -> str:
        match = re.search(r'```json\s*(.*?)\s*```', response_str, re.DOTALL)
        if match: return match.group(1).strip()
        start = response_str.find('{')
        end = response_str.rfind('}')
        if start != -1 and end != -1: return response_str[start:end+1]
        return response_str.strip()

    # ---------- JSON处理器（保持核心逻辑不变） ----------
    class JSONProc:
        # ... (This class remains unchanged) ...
        def __init__(self, parent: 'InfoExtractor'):
            self.parent = parent
        
        def merge_parameters(self, data: Dict[str, Any]) -> Dict[str, Any]:
            # ... (no changes in this method) ...
            merged_device_list = []
            stats = {'total_groups': 0, 'total_devices': 0, 'processed_devices': 0}
            if '设备列表' not in data or not isinstance(data['设备列表'], list):
                logger.warning("输入数据中缺少'设备列表'或格式不正确，无法执行合并。")
                return {"设备列表": merged_device_list, "stats": stats, "备注": data.get('备注')}
            stats['total_groups'] = len(data['设备列表'])
            for device_group in data['设备列表']:
                tag_nos = device_group.get('位号', [])
                common_params = device_group.get('标准化共用参数', {})
                diff_params = device_group.get('标准化不同参数', {})
                if not tag_nos: continue
                stats['total_devices'] += len(tag_nos)
                for tag_no in tag_nos:
                    params = common_params.copy()
                    for p_name, p_val_map in diff_params.items():
                        if isinstance(p_val_map, dict) and tag_no in p_val_map:
                            params[p_name] = p_val_map[tag_no]
                    merged_device_list.append({"位号": tag_no, "参数": params})
                    stats['processed_devices'] += 1
            logger.info(f"参数合并完成: {stats}")
            return {"设备列表": merged_device_list, "stats": stats, "备注": data.get('备注')}

        def extract_remarks(self, data: Dict[str, Any]) -> Optional[Dict[str, str]]:
            # ... (no changes in this method) ...
            remarks = data.get("备注")
            if isinstance(remarks, dict):
                return remarks
            return None
